from docx import Document
from .base import BaseConverter

class DocxConverter(BaseConverter):

    def convert(self, path: str) -> str:
        doc = Document(path)
        md_blocks = []

        for block in self._iter_blocks(doc):
            md_blocks.append(block)

        return "\n\n".join(md_blocks)

    # ---- Основная логика ----

    def _iter_blocks(self, doc):
        """Итерируем все блоки: параграфы и таблицы в порядке появления."""
        for element in doc.element.body:
            if element.tag.endswith('p'):
                yield self._paragraph_to_md(doc.paragraphs[self._p_index(doc, element)])
            elif element.tag.endswith('tbl'):
                table = doc.tables[self._t_index(doc, element)]
                yield self._table_to_md(table)

    def _p_index(self, doc, p_el):
        return [p._p for p in doc.paragraphs].index(p_el)

    def _t_index(self, doc, t_el):
        return [t._tbl for t in doc.tables].index(t_el)

    # ---- Преобразования ----

    def _paragraph_to_md(self, p):
        text = p.text.strip()
        if not text:
            return ""

        style = p.style.name if p.style else ""

        # Заголовки
        if style.startswith("Heading") or style.startswith("Заголовок"):
            try:
                level = int(''.join(ch for ch in style if ch.isdigit()))
            except:
                level = 1
            return f"{'#' * level} {text}"

        # Маркированные списки
        if "List Bullet" in style or "Список" in style:
            return f"- {text}"

        # Нумерованные списки
        if "List Number" in style or "Нумерованный" in style:
            return f"1. {text}"

        # Обычный параграф
        return text

    def _table_to_md(self, table):
        rows = list(table.rows)
        if not rows:
            return "_Пустая таблица_"

        md = []

        # читаем ячейки
        matrix = []
        for r in rows:
            matrix.append([self._clean_cell(c.text) for c in r.cells])

        headers = matrix[0]
        body = matrix[1:]

        # header
        md.append("| " + " | ".join(headers) + " |")
        md.append("| " + " | ".join("---" for _ in headers) + " |")

        # body
        for row in body:
            md.append("| " + " | ".join(row) + " |")

        return "\n".join(md)

    def _clean_cell(self, text):
        return text.replace("\n", " ").strip()
